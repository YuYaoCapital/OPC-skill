#!/usr/bin/env python3
"""横纵分析法深度研究：任相栋。"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime

OUTPUT_DIR = Path('/Users/r9/OPC/05_人力资源/员工档案')

def set_styles(doc):
    styles = doc.styles
    h1 = styles['Heading 1']
    h1.font.name = 'Microsoft YaHei'
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0, 51, 102)
    h1.font.bold = True
    
    h2 = styles['Heading 2']
    h2.font.name = 'Microsoft YaHei'
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0, 102, 153)
    h2.font.bold = True
    
    h3 = styles['Heading 3']
    h3.font.name = 'Microsoft YaHei'
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(0, 102, 153)
    h3.font.bold = True
    
    normal = styles['Normal']
    normal.font.name = 'SimSun'
    normal.font.size = Pt(11)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run.font.size = Pt(11)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(9)
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'SimSun'
                    run.font.size = Pt(9)
    doc.add_paragraph()

def generate(output_path=None):
    if output_path is None:
        output_path = OUTPUT_DIR / 'OPC_任相栋横纵分析法深度研究_v1.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    set_styles(doc)
    
    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('任相栋横纵分析法深度研究')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('从“交银新黄金一代”到“兴全日光基”再到清仓卸任')
    run.font.size = Pt(14)
    run.font.name = 'Microsoft YaHei'
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f'OPC 投研策略部  |  {datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(10)
    run.font.name = 'SimSun'
    
    doc.add_page_break()
    
    # 导语
    doc.add_heading('导语', level=1)
    doc.add_paragraph('任相栋的公募生涯，像极了一部“成也规模、败也规模”的教科书。')
    doc.add_paragraph('2015年，他在交银施罗德凭借交银先进制造一战成名，三年收益95.81%，年化21.7%，同类前5%，金牛奖拿到手软。2019年，他转会兴证全球，首发兴全合泰混合，单日认购超400亿，配售比例仅12%，成为当年最火爆的“日光基”。巅峰时期，他管理规模超过160亿，是兴全“五绝时代”之后最被看好的中生代。')
    doc.add_paragraph('但故事没有停在高光处。2022年发行的兴全合衡三年持有，成立至今仍亏损3.74%，最大回撤近46%。2025年末，管理规模缩水至83.91亿，退出百亿基金经理行列。2026年3月27日，任相栋因“个人原因”清仓式卸任，名下再无在管产品。')
    doc.add_paragraph('这不是一个简单的“明星陨落”故事。我们要问的是：任相栋的能力变了吗？他的投资框架失效了吗？规模、产品形态、市场风格、公司平台，哪个变量才是决定命运的关键？')
    
    # 第一部分：纵向分析
    doc.add_page_break()
    doc.add_heading('一、纵向分析：任相栋的十六年公募路', level=1)
    
    doc.add_heading('1. 交银时代（2010-2018）：制造业成长派的诞生', level=2)
    doc.add_paragraph('2010年，任相栋从上海交大金融硕士毕业后加入交银施罗德基金，从研究员做起，覆盖汽车、机械、国防军工等制造业。这个起点，几乎决定了他后来的全部投资底色。')
    doc.add_paragraph('制造业研究的训练是残酷的。它不像消费那样有漂亮的品牌故事，也不像科技那样有宏大的叙事空间。制造业研究员每天面对的是产能利用率、毛利率、订单周期、库存周转、资本开支。这种训练让任相栋养成了一种务实的选股习惯：看企业真实盈利能力，看行业竞争格局，看管理层执行能力。')
    doc.add_paragraph('2015年1月，他正式管理交银先进制造混合。那是一个特殊的时间点——A股刚经历完2014年底的金融地产暴动，创业板牛市进入中段。任相栋没有all in互联网+，而是继续在自己熟悉的制造业里精选个股。')
    doc.add_paragraph('结果证明，这种“老实”是对的。到2018年中卸任，交银先进制造任职回报95.81%，年化21.7%，同类排名前5%。更重要的是，这个业绩是在2015年6月股灾、2016年熔断、2018年贸易摩擦三轮大跌中跑出来的。')
    doc.add_paragraph('这个阶段的任相栋，有两个标签被确立下来：')
    add_bullet(doc, '制造业成长派：能力圈集中在汽车、机械、电子、电力设备')
    add_bullet(doc, '自下而上选股：不追热点，偏好左侧布局，强调安全边际')
    add_bullet(doc, '控制回撤意识：在2015年高位成立的产品中，回撤控制优于同类')
    
    doc.add_heading('2. 职业空窗期（2018-2019）：从金鹰到兴全', level=2)
    doc.add_paragraph('2018年离开交银后，任相栋短暂加入金鹰基金，任权益部上海权益总监，但没有管理公募产品。这段经历在公开资料中着墨不多，但它很可能是一个重要的观察窗口。')
    doc.add_paragraph('在金鹰的这段时间，任相栋从一个一线基金经理变成了一个“管理者”。他要搭建团队、培养研究员、制定投资流程。这种角色转换，对后来他在兴全的管理风格有一定影响。')
    doc.add_paragraph('2019年，他选择加盟兴证全球基金。这是一个非常理性的选择。兴全当时刚刚经历董承非、谢治宇等“五绝时代”核心人物的逐步淡出，急需一位能撑起门面、有业绩、有号召力、且风格与兴全“稳健成长”基因匹配的中生代。任相栋完美符合这个画像。')
    
    doc.add_heading('3. 兴全时代（2019-2026）：高光与困局', level=2)
    doc.add_heading('第一阶段：日光基与百亿规模（2019-2022）', level=3)
    doc.add_paragraph('2019年10月，兴全合泰混合首发。单日认购超400亿元，最终配售比例仅12%，创下当年发行纪录。这不仅是任相栋个人的高光，也是兴全在新老交替之际的一次成功品牌运作。')
    doc.add_paragraph('兴全合泰成立后的表现没有辜负这份热情。2020年初至2022年初，区间净值增长率88%。这个业绩在当时的市场环境下非常亮眼——它说明任相栋的制造业成长框架在新能源、汽车、电子的景气周期中依然有效。')
    doc.add_paragraph('2022年上半年，任相栋管理规模达到160.48亿元，跻身百亿基金经理行列。')
    
    doc.add_heading('第二阶段：合衡之困（2022-2025）', level=3)
    doc.add_paragraph('2022年1月，兴全合衡三年持有混合成立。这是任相栋职业生涯的转折点。')
    doc.add_paragraph('从产品设计的角度看，三年持有期本无问题。它可以锁定负债端，让基金经理做更长期的投资布局。但问题在于发行时点：2022年初，新能源、汽车、电子等成长赛道估值处于历史高位，市场风格即将切换。')
    doc.add_paragraph('兴全合衡成立后，市场经历了2022年的大跌、2023年的弱复苏、2024年的风格极致分化。任相栋擅长的制造业成长方向持续承压。到2026年3月任相栋卸任，兴全合衡任职回报为-3.74%，最大回撤近46%。')
    doc.add_paragraph('更麻烦的是三年持有期的流动性约束。投资者无法赎回，只能看着净值下跌。这种“套牢感”会迅速转化为对基金经理的信任崩塌。')
    
    doc.add_heading('第三阶段：卸任与转身（2025-2026）', level=3)
    doc.add_paragraph('2025年11月，兴全合泰增聘张传杰共同管理。2026年3月18日，兴全合泰再增聘谢书英，兴全合衡增聘张传杰。9天后，任相栋清仓式卸任。')
    doc.add_paragraph('这个时间线非常清晰：增聘→共管→卸任。这是公募行业明星基金经理离职的标准流程。')
    doc.add_paragraph('截至卸任，兴全合泰累计收益64.02%，仍然是一笔不错的长期回报；但兴全合衡的亏损和规模缩水，已经让他的公募生涯蒙上阴影。')
    
    doc.add_heading('4. 投资风格的稳定性与漂移', level=2)
    doc.add_paragraph('纵观任相栋的职业生涯，他的投资框架几乎没有发生本质变化：')
    add_bullet(doc, '自下而上精选个股')
    add_bullet(doc, '聚焦制造业成长方向')
    add_bullet(doc, '偏好左侧布局')
    add_bullet(doc, '强调安全边际')
    add_bullet(doc, '长期持有优质企业')
    doc.add_paragraph('这种稳定性，在顺风期是优势——投资者知道自己在买什么。但在逆风期，它变成了“路径依赖”。当市场风格从成长切换到价值、从制造切换到红利时，任相栋没有明显调整组合结构。')
    doc.add_paragraph('一个值得思考的问题是：这是“坚守能力圈”还是“固执”？')
    doc.add_paragraph('我的判断是：两者兼有。他的能力圈确实在制造业，切换方向未必能做好；但面对46%的最大回撤，适当的仓位管理和行业分散是必要的。兴全合衡的问题，不完全来自选股，也来自对系统性风险的低估。')
    
    # 第二部分：横向分析
    doc.add_page_break()
    doc.add_heading('二、横向分析：任相栋与同类基金经理对比', level=1)
    
    doc.add_heading('1. 与“交银新黄金一代”对比', level=2)
    doc.add_paragraph('2015年前后，交银施罗德涌现出一批优秀的中生代基金经理，被市场称为“新黄金一代”。代表人物包括何帅、杨浩、任相栋等。')
    
    add_table(doc, ['维度', '任相栋', '何帅', '杨浩'], [
        ['核心能力圈', '制造业（汽车、机械、电子、电力设备）', '医药、消费、TMT', 'TMT、消费、制造业'],
        ['投资风格', '成长价值均衡，偏左侧', '深度价值成长，逆向投资', '成长风格，行业轮动'],
        ['代表作业绩', '交银先进制造3年+95.81%', '交银优势行业长期优秀', '交银定期支付双息平衡长期优秀'],
        ['职业路径', '交银→金鹰→兴全→卸任/奔私', '留守交银，逐步成为核心', '留守交银，后遇业绩压力'],
        ['核心标签', '制造业成长派', '逆向价值派', '均衡成长派'],
    ])
    
    doc.add_paragraph('对比之下，任相栋是最早离开交银的一位。他选择了加盟更大的平台、管理更大的规模，但也因此承受了更大的压力。何帅和杨浩留守交银，虽然后来也各有起伏，但至少保持了平台稳定性和投资连贯性。')
    doc.add_paragraph('这个对比揭示了一个残酷的现实：基金经理的职业生涯，不只是投资能力的比拼，也是平台选择、时机把握、规模管理的综合结果。')
    
    doc.add_heading('2. 与兴全系基金经理对比', level=2)
    doc.add_paragraph('兴证全球基金以“稳健、均衡、长期”著称。任相栋在兴全的六年，可以放在这个坐标系中观察。')
    
    add_table(doc, ['维度', '任相栋', '谢治宇', '董承非', '乔迁'], [
        ['兴全角色', '中生代旗帜', '五绝时代核心', '五绝时代核心', '中生代代表'],
        ['投资风格', '成长制造，偏左侧', '均衡成长', '均衡价值', '均衡成长'],
        ['代表作', '兴全合泰', '兴全合润', '兴全趋势', '兴全商业模式'],
        ['回撤控制', '一般（合衡最大回撤近46%）', '较好', '优秀', '较好'],
        ['规模管理', '巅峰160亿，卸任前84亿', '长期200亿+', '长期百亿以上', '百亿左右'],
    ])
    
    doc.add_paragraph('兴全的核心竞争力，从来不是某一位明星基金经理的极致业绩，而是整个投研体系对风险和收益的均衡把握。任相栋的加入，本是为了承接兴全的权益大旗，但他的风格比传统兴全产品更偏成长、更集中、波动更大。')
    doc.add_paragraph('兴全合泰在顺风期跑出了88%的区间收益，这让它看起来像是一只“加强版兴全产品”；但兴全合衡在逆风期的46%回撤，又暴露了它与兴全“稳健”基因的距离。')
    doc.add_paragraph('一个可能的结论是：任相栋从来不是一个典型的“兴全系”基金经理。他是被兴全平台放大了的“任相栋”。')
    
    doc.add_heading('3. 与全市场均衡成长型基金经理对比', level=2)
    doc.add_paragraph('如果把任相栋放到全市场同类基金经理中比较，他的特点会更清晰。')
    
    doc.add_paragraph('第一，他的选股能力是被验证过的。交银先进制造、兴全合泰的长期业绩，都说明他在制造业领域有持续的阿尔法。')
    doc.add_paragraph('第二，他的能力圈相对集中。相比朱少醒、傅鹏博、周蔚文这些“全市场选股”的老将，任相栋的覆盖范围明显更窄。这在顺风期是优势，在逆风期是劣势。')
    doc.add_paragraph('第三，他的规模管理能力存在短板。160亿的管理规模，对一位偏自下而上选股的基金经理来说，已经是一种负担。很多优质中小盘制造业公司无法容纳这么大的资金量，迫使他向大市值龙头集中，削弱了阿尔法的来源。')
    doc.add_paragraph('第四，他在产品形态选择上吃亏。三年持有期本是为了长期投资，但在市场高点发行、且组合风格与市场风向背离的情况下，流动性锁死反而放大了持有人的痛苦。')
    
    # 第三部分：横纵交汇
    doc.add_page_break()
    doc.add_heading('三、横纵交汇：任相栋的公募生涯启示', level=1)
    
    doc.add_heading('1. 历史如何塑造了今天的位置', level=2)
    doc.add_paragraph('任相栋今天的位置，是早期一系列成功决策的累积结果，也是这些决策在特定市场环境下的反噬。')
    doc.add_paragraph('他的制造业背景、自下而上选股、左侧布局，这些在2015-2021年的成长股牛市中是无往不利的武器。但2022年之后，市场风格转向价值、红利、低估值，他的武器库里没有对应的工具。')
    doc.add_paragraph('他没有像一些基金经理那样切换风格，这可能源于两个原因：一是真正的能力圈自信，二是管理规模太大、调仓成本太高。无论哪个原因，结果都是他在逆风期暴露得比别人更彻底。')
    
    doc.add_heading('2. 优势的历史根源与潜在包袱', level=2)
    doc.add_paragraph('今天的优势：')
    add_bullet(doc, '制造业选股阿尔法：来自交银时代的研究积累')
    add_bullet(doc, '长期视角：愿意陪伴企业成长，不频繁换手')
    add_bullet(doc, '平台号召力：兴全日光基证明了他在渠道端的品牌价值')
    
    doc.add_paragraph('今天的包袱：')
    add_bullet(doc, '规模诅咒：160亿规模迫使他放弃部分中小盘机会')
    add_bullet(doc, '路径依赖：能力圈过于集中在制造业成长方向')
    add_bullet(doc, '产品形态约束：三年持有期在错误时点发行')
    add_bullet(doc, '平台期待落差：兴全需要的是“稳健成长”，但他更偏“锐利成长”')
    
    doc.add_heading('3. 三个未来剧本', level=2)
    
    doc.add_heading('最可能的剧本：奔私后做小规模精品基金', level=3)
    doc.add_paragraph('任相栋下一站大概率是私募。在私募领域，他可以摆脱公募的排名压力和规模约束，管理几十亿规模，回归自己最擅长的制造业成长选股。如果他能在私募中控制规模、精选个股，依然有机会复制交银先进制造时期的辉煌。')
    
    doc.add_heading('最危险的剧本：规模再次膨胀，重蹈覆辙', level=3)
    doc.add_paragraph('如果奔私后产品发行火爆，规模迅速突破百亿，他可能会再次陷入同样的困境。私募虽然比公募灵活，但规模仍然是阿尔法的敌人。如果他不能主动限制规模，历史很可能重演。')
    
    doc.add_heading('最乐观的剧本：风格回归，王者归来', level=3)
    doc.add_paragraph('如果未来1-2年市场风格重新转向成长制造，任相栋的投资框架将再次大放异彩。他的能力圈没有退化，只是暂时与市场风格不匹配。一旦风来了，他仍然是那个能跑出阿尔法的基金经理。')
    
    doc.add_heading('4. 对 OPC 的启示', level=2)
    doc.add_paragraph('任相栋的案例，对 OPC 这种买方投顾机构有几个重要启示：')
    add_bullet(doc, '基金经理的能力圈不是无限的，投顾方案要匹配基金经理的真实风格')
    add_bullet(doc, '规模是业绩的敌人，百亿以上要主动警惕阿尔法衰减')
    add_bullet(doc, '产品形态很重要，三年持有期不是万能的，高点发行的锁定期会放大客户不满')
    add_bullet(doc, '不要把平台品牌与个人品牌混为一谈，兴全合泰的成功不等于任相栋所有产品都会成功')
    add_bullet(doc, '对客户的预期管理，要基于基金经理的能力圈和当前市场风格，而不是历史业绩')
    
    doc.add_heading('结语', level=1)
    doc.add_paragraph('任相栋不是能力不行，而是在错误的时间、用错误的产品形态、管了过大的规模、坚持了与市场风格暂时不匹配的方向。')
    doc.add_paragraph('他的故事提醒我们：评价一位基金经理，不能只看高光时刻的收益，还要看他在逆风期的应对、他对规模的敬畏、他对能力圈的诚实。')
    doc.add_paragraph('对 OPC 来说，任相栋是一个值得长期跟踪的研究对象。如果未来他在私募中控制规模、回归制造业选股，很可能是一个不错的配置机会。但在那之前，我们需要保持观察。')
    
    # 免责声明
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('免责声明：本报告基于公开信息和历史数据整理，不构成投资建议。基金经理历史业绩不代表未来表现，投资有风险，入市需谨慎。')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.name = 'SimSun'
    
    doc.save(str(output_path))
    return output_path

if __name__ == '__main__':
    path = generate()
    print(f"任相栋横纵分析 Word 已生成：{path}")
