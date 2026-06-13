#!/usr/bin/env python3
"""生成面向基金代销机构理财经理的 OPC Skill 使用介绍 Word。"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime

OUTPUT_DIR = Path('/Users/r9/OPC/02_投顾服务')

def set_styles(doc):
    styles = doc.styles
    h1 = styles['Heading 1']
    h1.font.name = 'Microsoft YaHei'
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0, 51, 102)
    h1.font.bold = True
    
    h2 = styles['Heading 2']
    h2.font.name = 'Microsoft YaHei'
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0, 102, 153)
    h2.font.bold = True
    
    normal = styles['Normal']
    normal.font.name = 'SimSun'
    normal.font.size = Pt(11)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run.font.size = Pt(11)

def add_example(doc, title, steps):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    for step in steps:
        add_bullet(doc, step)

def generate(output_path=None):
    if output_path is None:
        output_path = OUTPUT_DIR / 'OPC_Skill_理财经理使用指南_v1.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    set_styles(doc)
    
    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('OPC Skill')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('理财经理使用指南')
    run.font.size = Pt(20)
    run.font.name = 'Microsoft YaHei'
    
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('一个人，也能拥有投研团队的专业能力')
    run.font.size = Pt(12)
    run.font.name = 'SimSun'
    
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run(f'版本：v1.0  |  日期：{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(10)
    run.font.name = 'SimSun'
    
    doc.add_page_break()
    
    # 1. 这是给谁看的
    doc.add_heading('一、这份指南是写给谁的？', level=1)
    doc.add_paragraph('如果你是基金代销机构的理财经理，每天面临这些场景，这份指南就是为你写的：')
    add_bullet(doc, '客户问你：“这只基金怎么样？”你只知道看近一年收益，讲不出深度')
    add_bullet(doc, '客户亏钱了，你想安抚却不知道说什么专业又有温度')
    add_bullet(doc, '银行要你做基金沙龙，你熬夜做 PPT，内容还是从网上拼凑的')
    add_bullet(doc, '领导让你写周报月报，你花了半天整理数据，排版还不专业')
    add_bullet(doc, '你想给客户做组合诊断，但Excel模型不会搭，wind也没有')
    
    doc.add_paragraph('OPC Skill 相当于给你配了一个“24 小时在线的投研 + 投顾助理团队”。你不用懂编程，不用买 wind，只要会说话，它就能帮你出报告、写话术、做分析。')
    
    # 2. OPC 是什么
    doc.add_heading('二、OPC 是什么？一句话说明白', level=1)
    doc.add_paragraph('OPC 是一套“AI 投顾工具包”，把专业投研团队的认知和能力，封装成一个个 Skill（技能包）。')
    doc.add_paragraph('你可以把它理解成：你的手机里多了一个“投顾大脑”，你想让它干嘛，直接说人话就行。')
    
    doc.add_heading('打个比方', level=2)
    doc.add_paragraph('以前你要写一篇基金分析报告，需要：')
    add_bullet(doc, '自己查数据')
    add_bullet(doc, '自己写框架')
    add_bullet(doc, '自己排版')
    add_bullet(doc, '自己检查有没有错')
    doc.add_paragraph('用了 OPC Skill 之后，你只需要说：')
    p = doc.add_paragraph()
    run = p.add_run('“帮我写一份某某基金的深度分析报告，要包含业绩归因、持仓分析、基金经理评价。”')
    run.italic = True
    run.font.name = 'SimSun'
    doc.add_paragraph('剩下的它自动帮你完成。')
    
    # 3. 你能用它做什么
    doc.add_page_break()
    doc.add_heading('三、理财经理能用它做什么？', level=1)
    
    doc.add_heading('1. 基金深度评价', level=2)
    add_example(doc, '你说：', [
        '“评价一下这只固收+基金，值不值得推荐给稳健型客户？”',
        '“分析一下这只科技主题基金的风险收益特征。”'
    ])
    doc.add_paragraph('它能输出：收益指标、风险指标、最大回撤恢复天数、收益归因、调仓有效性、基金经理画像、适合什么客户类型。')
    
    doc.add_heading('2. 客户陪伴话术', level=2)
    add_example(doc, '你说：', [
        '“客户持有的权益基金回撤了 15%，怎么安抚？”',
        '“给我一份市场波动时的客户维护话术。”'
    ])
    doc.add_paragraph('它能输出：差异化安抚话术（按客户类型、盈亏状态、超配情况区分），结合最新市场分析和重仓股解读。')
    
    doc.add_heading('3. 组合诊断与调仓建议', level=2)
    add_example(doc, '你说：', [
        '“诊断一下这个客户的基金组合。”',
        '“这个组合权益超配了，帮我写一份调仓发车文案。”'
    ])
    doc.add_paragraph('它能输出：组合风险收益分析、资产配置偏离度、再平衡建议、调仓理由和客户可理解的沟通文案。')
    
    doc.add_heading('4. 研究报告与公众号内容', level=2)
    add_example(doc, '你说：', [
        '“帮我写一篇关于风格切换的公众号文章。”',
        '“点评一下宁德时代最新季报。”',
        '“研究一下半导体行业，输出一份行业报告。”'
    ])
    doc.add_paragraph('它能输出：专业研报、公众号长文、财经热点图、晨会纪要等多种格式。')
    
    doc.add_heading('5. 日常运营材料', level=2)
    add_example(doc, '你说：', [
        '“生成今日市场点评。”',
        '“整理一份本周工作成果汇报。”',
        '“做一个客户年度回顾报告。”'
    ])
    doc.add_paragraph('它能输出：Word、Excel、PPT、PDF 等多种格式的专业材料。')
    
    # 4. 具体怎么用
    doc.add_page_break()
    doc.add_heading('四、具体怎么用？三步上手', level=1)
    
    doc.add_heading('第一步：安装 OPC Skill', level=2)
    doc.add_paragraph('把 OPC 提供的 Skill 压缩包解压到你的 Kimi Code CLI 技能目录里。')
    code = doc.add_paragraph()
    run = code.add_run('cd ~/.kimi/skills\nunzip /path/to/OPC_skills_20260613.zip')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    doc.add_paragraph('如果你不会操作，让 IT 同事帮你放一下，只要放到 .kimi/skills/ 目录下就行。')
    
    doc.add_heading('第二步：打开 Kimi Code CLI', level=2)
    doc.add_paragraph('就像打开一个智能助手对话框。不需要写代码，直接打字说话。')
    
    doc.add_heading('第三步：说人话，让它干活', level=2)
    doc.add_paragraph('下面是一些真实可用的示例，你可以直接复制粘贴进去：')
    
    examples = [
        ('基金研究', '“帮我深度研究一下易方达蓝筹精选，看看适不适合推荐给高净值客户。”'),
        ('客户安抚', '“客户买的偏股基金亏了 10%，帮我写一段专业又有温度的安抚话术。”'),
        ('组合诊断', '“这是我的客户持仓Excel，帮我诊断一下组合有没有问题。”'),
        ('市场点评', '“今天成长风格大跌，帮我写一段收盘点评，安抚客户情绪。”'),
        ('行业研究', '“帮我研究一下 AI 算力产业链，输出一份 10 页以内的报告。”'),
        ('内容生产', '“帮我写一篇关于红利策略的公众号文章，要通俗易懂。”'),
    ]
    
    for title, text in examples:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}：')
        run.bold = True
        run.font.name = 'Microsoft YaHei'
        run = p.add_run(text)
        run.font.name = 'SimSun'
    
    # 5. 使用技巧
    doc.add_page_break()
    doc.add_heading('五、说得更清楚，结果就更好', level=1)
    doc.add_paragraph('OPC Skill 很依赖你说的话。说得越具体，输出越贴近你的需求。')
    
    doc.add_heading('好的提问方式', level=2)
    add_bullet(doc, '“评价这只基金” → 太笼统')
    add_bullet(doc, '“评价这只基金，客户是 C3 稳健型，持有期 1-3 年” → 更好')
    add_bullet(doc, '“写一段安抚话术” → 太笼统')
    add_bullet(doc, '“客户 55 岁，买了 50 万权益基金，现在浮亏 12%，写一段 200 字的安抚话术” → 更好')
    
    doc.add_heading('多给上下文', level=2)
    add_bullet(doc, '客户年龄、风险偏好、投资目标')
    add_bullet(doc, '基金代码、持有金额、买入时间')
    add_bullet(doc, '你想要什么格式：Word / PPT / PDF / Excel')
    add_bullet(doc, '你想要什么风格：专业严谨 / 通俗易懂 / 轻松幽默')
    
    # 6. 常见误区
    doc.add_heading('六、使用前你要知道的几件事', level=1)
    add_bullet(doc, 'OPC Skill 是“助手”，不是“股神”。它帮你提高效率、规范输出，但投资决策仍需你独立判断')
    add_bullet(doc, '它不能直接帮客户下单，不能代替合规审核')
    add_bullet(doc, '给客户看的材料，建议你先过一遍，加上自己的理解和客户化表达')
    add_bullet(doc, '部分功能需要本地数据（如客户持仓 Excel），没有数据时它只能给通用模板')
    add_bullet(doc, '它依赖 Kimi Code CLI 运行环境，字体和 Python 包需要提前配置好')
    
    # 7. 总结
    doc.add_heading('七、一句话总结', level=1)
    p = doc.add_paragraph()
    run = p.add_run('OPC Skill = 一个随时待命的投研 + 投顾内容团队。')
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Microsoft YaHei'
    doc.add_paragraph('你负责和客户建立信任、理解需求、做最终判断；它负责查数据、写报告、做分析、出话术。')
    doc.add_paragraph('一个人，也能做出一个团队的专业感。')
    
    doc.save(str(output_path))
    return output_path

if __name__ == '__main__':
    path = generate()
    print(f"理财经理指南 Word 已生成：{path}")
