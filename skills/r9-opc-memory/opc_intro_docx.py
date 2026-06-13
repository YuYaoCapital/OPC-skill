#!/usr/bin/env python3
"""生成 OPC 投顾公司对外宣传介绍 Word。"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime

OUTPUT_DIR = Path('/Users/r9/OPC/00_公司治理')

def set_heading_style(doc):
    """设置标题样式。"""
    styles = doc.styles
    
    # 标题 1
    h1 = styles['Heading 1']
    h1.font.name = 'Microsoft YaHei'
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0, 51, 102)
    h1.font.bold = True
    
    # 标题 2
    h2 = styles['Heading 2']
    h2.font.name = 'Microsoft YaHei'
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0, 51, 102)
    h2.font.bold = True
    
    # 正文
    normal = styles['Normal']
    normal.font.name = 'SimSun'
    normal.font.size = Pt(11)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run.font.size = Pt(11)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(10)
    
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'SimSun'
                    run.font.size = Pt(10)
    
    doc.add_paragraph()

def generate(output_path=None):
    if output_path is None:
        output_path = OUTPUT_DIR / 'OPC_公司介绍与Skill体系_v1.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    set_heading_style(doc)
    
    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('OPC 投顾公司')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('公司介绍与 AI Skill 体系')
    run.font.size = Pt(18)
    run.font.name = 'Microsoft YaHei'
    
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('以认知产品化为核心的买方投顾机构')
    run.font.size = Pt(12)
    run.font.name = 'SimSun'
    
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run(f'版本：v1.0  |  日期：{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(10)
    run.font.name = 'SimSun'
    
    doc.add_page_break()
    
    # 1. 公司定位
    doc.add_heading('一、公司定位', level=1)
    doc.add_paragraph('OPC（One-Person Company）投顾公司是一家以“认知产品化”为核心的买方投顾机构。')
    doc.add_paragraph('我们坚信：在 AI 时代，投顾服务的核心竞争力不是规模，而是“可复制的深度认知”。OPC 通过董事长 R9 的专业判断 + 一套完整的 Agent 体系，把研究、服务、运营、合规全流程标准化，实现“一人服务千人”的买方投顾模式。')
    
    doc.add_heading('核心特征', level=2)
    add_bullet(doc, '买方立场：收投顾费，不卖产品，不与销售佣金挂钩')
    add_bullet(doc, 'AI 原生：投研、投顾、运营、合规均由专业 Agent 协同完成')
    add_bullet(doc, '认知产品化：研究成果封装为可复用、可订阅的 AI Skill')
    add_bullet(doc, 'ToC 为主、ToB 为辅：ToC 买方投顾收入占比 ≥70%，ToB Skill 会员 ≤30%')
    
    # 2. 商业模式
    doc.add_heading('二、商业模式', level=1)
    doc.add_paragraph('OPC 采用“双轮驱动”模式：')
    
    doc.add_heading('ToC 买方投顾', level=2)
    add_bullet(doc, '服务对象：有认知但没时间的中产专业人士')
    add_bullet(doc, '服务内容：资产配置、组合诊断、再平衡、投后陪伴')
    add_bullet(doc, '收费模式：按 AUM 收取年度投顾费')
    add_bullet(doc, '服务规模：当前 87 位客户，AUM 约 3000 万，边际服务成本趋近于零')
    
    doc.add_heading('ToB AI Skill 会员', level=2)
    add_bullet(doc, '输出对象：银行、券商、理财子、财富管理机构')
    add_bullet(doc, '输出内容：基金评价模型、估值工具、客户话术库、研究报告模板')
    add_bullet(doc, '收费模式：会员订阅 / 单模块授权')
    add_bullet(doc, '战略约束：ToB 收入 ≤30%，不做定制开发，R9 每月 ToB 投入 ≤2 天')
    
    # 3. 组织架构
    doc.add_page_break()
    doc.add_heading('三、组织架构', level=1)
    doc.add_paragraph('OPC 采用“1 人 + Agent 体系”的虚拟公司架构：')
    
    org_rows = [
        ['董事长 / 投决会主任', 'R9', '最终决策权、一票否决权'],
        ['CEO / 投决会副主任 / 超级路由', 'Luce', '统筹全局、任务分发、投决会运作'],
        ['投研策略部总经理', 'Atlas', '投资研究、策略开发、资产配置、认知产品化'],
        ['投顾服务部总经理', 'Mira', '客户获客、投顾签约、客户陪伴、ToC 收费服务'],
        ['交易运营部总经理', 'Vega', '交易执行、运营支持、技术基建、底层费用管理'],
        ['合规风控部总经理', 'Sage', '合规审查、风险控制、法务支持、内控审计'],
    ]
    add_table(doc, ['角色', '负责人', '核心职能'], org_rows)
    
    doc.add_heading('投研策略部研究员', level=2)
    add_bullet(doc, 'Helios（大类资产研究员）：SAA/TAA、风格轮动、资产周期定位')
    add_bullet(doc, 'Terra（宏观研究员）：宏观经济、政策解读、全球宏观环境研判')
    add_bullet(doc, 'Mercury（行业研究员）：A 股/港股行业景气度、产业链研究')
    add_bullet(doc, 'Castor（基金评价研究员）：R9Alpha 基金评价模型、基金经理画像')
    add_bullet(doc, 'Orion（组合研究员）：组合诊断、再平衡、绩效归因')
    
    doc.add_heading('投顾服务部团队', level=2)
    add_bullet(doc, 'Aiden（顾问策略师）：服务产品设计、客户旅程、SOP')
    add_bullet(doc, 'Belle（内容运营官）：内容策略、多形态内容生产、IP 运营')
    add_bullet(doc, 'Clara（客户成功经理）：客户分层运营、续约增购、NPS 管理')
    add_bullet(doc, 'Dylan（投顾交付主管）：投顾方案一线交付、客户沟通、投后跟踪')
    
    # 4. Skill 体系
    doc.add_page_break()
    doc.add_heading('四、AI Skill 体系', level=1)
    doc.add_paragraph('OPC 的全部专业能力已被封装为 17 个 Kimi Skill，覆盖投研、投顾、运营、合规、记忆五大领域。这些 Skill 让 OPC 的认知能力可以被复用、被分发、被订阅。')
    
    skill_rows = [
        ['r9-opc-ceo', 'CEO 统筹', 'Luce 全局调度、投决会运作'],
        ['r9-opc-research', '投研统筹', 'Atlas 研究任务分配与投决会支持'],
        ['r9-opc-research-asset', '大类资产', 'Helios SAA/TAA 研究'],
        ['r9-opc-research-macro', '宏观研究', 'Terra 宏观跟踪与政策解读'],
        ['r9-opc-research-sector', '行业研究', 'Mercury 行业景气度与产业链'],
        ['r9-opc-research-fund', '基金评价', 'Castor R9Alpha 基金评价'],
        ['r9-opc-research-portfolio', '组合研究', 'Orion 组合诊断与再平衡'],
        ['r9-opc-advisory', '投顾统筹', 'Mira 投顾服务统筹'],
        ['r9-opc-advisory-strategy', '顾问策略', 'Aiden 服务产品设计'],
        ['r9-opc-advisory-content', '内容运营', 'Belle 内容生产与 IP 运营'],
        ['r9-opc-advisory-success', '客户成功', 'Clara 客户分层与 NPS'],
        ['r9-opc-advisory-delivery', '投顾交付', 'Dylan 一线交付与陪伴'],
        ['r9-opc-operations', '交易运营', 'Vega 数据、交易、技术基建'],
        ['r9-opc-compliance', '合规风控', 'Sage 合规审查与风险'],
        ['r9-opc-memory', '公司记忆', '会议纪要、指令跟踪、历史检索'],
        ['r9-workbench', '超级路由', '统一调度全部 Skill'],
        ['china-market-data', '数据接口', 'A 股/港股/基金/宏观数据'],
    ]
    add_table(doc, ['Skill 名称', '所属领域', '主要职责'], skill_rows)
    
    # 5. 数据与技术
    doc.add_page_break()
    doc.add_heading('五、数据与技术基建', level=1)
    doc.add_heading('数据接入', level=2)
    add_bullet(doc, '市场数据：AKShare、Tushare、Baostock、Wind/iFinD（专业终端）')
    add_bullet(doc, '客户数据：持仓、交易、收益、风险测评 Excel/API 接入')
    add_bullet(doc, '当前成熟度：L2 半自动化，目标 L3 全自动化')
    
    doc.add_heading('技术栈', level=2)
    add_bullet(doc, 'Agent 调度：Kimi Code CLI + r9-workbench 超级路由')
    add_bullet(doc, '报告生成：Python + fpdf2 / python-docx / openpyxl')
    add_bullet(doc, '数据存储：PostgreSQL（迁移中）+ 自动备份')
    add_bullet(doc, '知识管理：本地文件归档 + JSON 索引 + r9-opc-memory')
    
    # 6. 合规
    doc.add_heading('六、合规与资质', level=1)
    doc.add_paragraph('OPC 在取得自有投顾牌照前，采用“与持牌机构合作”模式，由持牌机构作为面向客户提供投顾服务的主体，OPC 定位为策略内容供应商与技术服务方。')
    add_bullet(doc, 'R9 需在 3 个月内取得证券从业资格')
    add_bullet(doc, '6 个月内取得证券投资咨询执业资格')
    add_bullet(doc, '所有对外输出经 Sage 合规审核并留痕')
    add_bullet(doc, '客户适当性管理：C1-C5 风险等级与 L0-L4 服务等级匹配')
    
    # 7. 成果里程碑
    doc.add_heading('七、建设成果', level=1)
    add_bullet(doc, '2026-06-11：公司治理架构确立，四部门 Agent 体系建立')
    add_bullet(doc, '2026-06-11：首次投决会召开，形成 6 项决议 + 12 项行动项')
    add_bullet(doc, '2026-06-12：完成投顾业务合规框架、客户数据接入测试、客户分层方案')
    add_bullet(doc, '2026-06-13：完成 OPC 公司记忆系统（r9-opc-memory），实现会议纪要、指令跟踪、历史检索')
    add_bullet(doc, '持续输出：每日资本市场简报、行业专题研究、基金评价报告、客户陪伴内容')
    
    # 8. Skill 使用说明
    doc.add_page_break()
    doc.add_heading('八、如何获取与使用 OPC Skill', level=1)
    doc.add_paragraph('OPC 的 Skill 体系基于 Kimi Code CLI 的 Skill 机制构建。每个 Skill 是一个独立目录，包含 SKILL.md 说明文件和可选的脚本/资源。')
    
    doc.add_heading('1. 获取 Skill', level=2)
    add_bullet(doc, '方式一：接收 OPC 提供的 Skill 压缩包（如 OPC_skills_20260613.zip）')
    add_bullet(doc, '方式二：从 OPC 的代码仓库或共享目录复制单个 Skill 目录')
    
    doc.add_heading('2. 安装 Skill', level=2)
    doc.add_paragraph('以 macOS / Linux 为例，将 Skill 目录解压或复制到用户主目录下的 .kimi/skills/：')
    code = doc.add_paragraph()
    run = code.add_run('cd ~/.kimi/skills\n# 如果是 zip 包\nunzip /path/to/OPC_skills_20260613.zip\n\n# 如果是单个 Skill\ncp -r /path/to/r9-opc-memory ~/.kimi/skills/')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_paragraph('安装后的目录结构示例：')
    code2 = doc.add_paragraph()
    run2 = code2.add_run('~/.kimi/skills/\n├── r9-workbench/\n│   └── SKILL.md\n├── r9-opc-memory/\n│   ├── SKILL.md\n│   ├── archive.py\n│   └── ...\n└── china-market-data/\n    └── SKILL.md')
    run2.font.name = 'Courier New'
    run2.font.size = Pt(9)
    
    doc.add_heading('3. 使用 Skill', level=2)
    add_bullet(doc, '打开 Kimi Code CLI 或支持的代码助手')
    add_bullet(doc, '直接说人话触发 Skill，例如：“帮我研究一下贵州茅台”会自动调用 r9-workbench 路由')
    add_bullet(doc, 'Skill 的触发词在 SKILL.md 的 description 中定义，系统会自动匹配')
    add_bullet(doc, '对于带脚本的 Skill（如 r9-opc-memory），Agent 会根据 SKILL.md 指引执行对应 Python 脚本')
    
    doc.add_heading('4. 注意事项', level=2)
    add_bullet(doc, '每个 Skill 目录必须包含 SKILL.md，否则不会被识别')
    add_bullet(doc, '部分 Skill 依赖本地字体（Songti.ttc、STHeiti Light.ttc）和 Python 包（fpdf2、openpyxl 等）')
    add_bullet(doc, 'OPC Skill 中的路径（如 /Users/r9/OPC/）需根据使用者的实际环境调整')
    add_bullet(doc, '对外分发时，建议移除内部记忆索引（memory_index.json）和 tracker_data.json 等私有数据')
    
    doc.save(str(output_path))
    return output_path

if __name__ == '__main__':
    path = generate()
    print(f"公司介绍 Word 已生成：{path}")
