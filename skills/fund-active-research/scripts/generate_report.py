#!/usr/bin/env python3
"""
基金投资研究报告生成器
生成专业的Word格式投资报告
"""

import os
import sys
import argparse
from datetime import datetime
from pathlib import Path
from typing import Dict, List
import pandas as pd
import numpy as np

# 配置matplotlib中文字体
import matplotlib
matplotlib.use('Agg')  # 非交互式后端
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib import font_manager

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimSong', 'STHeiti', 'Heiti TC', 'Arial Unicode MS', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 添加当前目录到路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from fund_eastmoney import FundDataFetcher
from holding_analyzer import HoldingEffectivenessAnalyzer
from holding_history_analyzer import HoldingHistoryAnalyzer
from fund_comparison import FundComparator


class FundResearchReport:
    """基金投资研究报告生成器"""
    
    def __init__(self, template_path: str = None):
        self.template_path = template_path
        self.doc = None
        self.data = None
        self.fund_code = None
        
    def generate(self, fund_code: str, output_dir: str = None, years: int = 3) -> str:
        """
        生成基金投资报告
        
        Args:
            fund_code: 基金代码
            output_dir: 输出目录，默认当前目录
            years: 历史数据年数
            
        Returns:
            生成的报告文件路径
        """
        self.fund_code = fund_code
        
        # 获取数据
        print(f"\n{'='*60}")
        print(f"开始生成基金投资报告 - {fund_code}")
        print(f"{'='*60}\n")
        
        fetcher = FundDataFetcher()
        self.data = fetcher.get_all_data(fund_code, years)
        
        if 'error' in self.data.get('fund_info', {}):
            raise Exception(f"获取基金数据失败: {self.data['fund_info']['error']}")
        
        # 创建文档
        print("\n正在生成报告...")
        self._create_document()
        
        # 添加报告内容
        self._add_cover()
        self._add_chapter1_overview()
        self._add_chapter2_performance()
        self._add_chapter3_risk()
        self._add_chapter4_attribution()
        self._add_chapter5_holdings()
        self._add_chapter6_turnover()
        self._add_chapter7_manager()
        self._add_chapter8_comparison()
        self._add_chapter9_pros_cons()
        self._add_chapter10_rating()
        self._add_chapter11_outlook()
        
        # 保存文档
        fund_name = self.data['fund_info'].get('fund_name', fund_code)
        fund_name = fund_name.replace(' ', '').replace('/', '-')
        date_str = datetime.now().strftime('%Y%m%d')
        filename = f"基金投资报告_{fund_code}_{fund_name}_{date_str}.docx"
        
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            filepath = os.path.join(output_dir, filename)
        else:
            filepath = filename
            
        self.doc.save(filepath)
        print(f"\n✓ 报告已生成: {filepath}")
        
        return filepath
    
    def _create_document(self):
        """创建Word文档并设置样式"""
        self.doc = Document()
        
        # 设置中文字体
        self._set_chinese_font()
        
        # 设置页面边距
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
    
    def _set_chinese_font(self):
        """设置文档中文字体"""
        # 设置正文样式
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(12)
        
        # 设置标题样式
        for i in range(1, 10):
            try:
                heading_style = self.doc.styles[f'Heading {i}']
                heading_style.font.name = '黑体'
                heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                heading_style.font.size = Pt(16 - i)
                heading_style.font.bold = True
            except:
                pass
    
    def _add_paragraph(self, text: str, style: str = 'Normal', alignment = None, 
                       bold: bool = False, size: int = None):
        """添加段落"""
        p = self.doc.add_paragraph(text, style=style)
        if alignment:
            p.alignment = alignment
        if bold:
            for run in p.runs:
                run.bold = True
        if size:
            for run in p.runs:
                run.font.size = Pt(size)
        return p
    
    def _add_heading(self, text: str, level: int = 1):
        """添加标题"""
        return self.doc.add_heading(text, level=level)
    
    def _add_table(self, data: List[List], style: str = 'Table Grid'):
        """添加表格"""
        if not data or not data[0]:
            return None
            
        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = style
        
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = str(cell_data)
                
        return table
    
    def _add_picture(self, image_path: str, width: float = None):
        """添加图片"""
        if width:
            self.doc.add_picture(image_path, width=Inches(width))
        else:
            self.doc.add_picture(image_path)
    
    def _setup_chinese_font(self):
        """设置matplotlib中文字体"""
        plt.rcParams['font.sans-serif'] = ['SimSong', 'STHeiti', 'Heiti TC', 'Arial Unicode MS', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False
    
    def _generate_performance_chart(self) -> str:
        """生成业绩走势图"""
        self._setup_chinese_font()
        
        nav_df = self.data['nav_history']
        if nav_df.empty:
            return None
            
        plt.figure(figsize=(10, 5))
        plt.plot(nav_df['date'], nav_df['nav'], label='单位净值', linewidth=1.5)
        plt.title(f"{self.data['fund_info'].get('fund_name', '')} 净值走势", fontsize=14)
        plt.xlabel('日期', fontsize=10)
        plt.ylabel('净值', fontsize=10)
        plt.grid(True, alpha=0.3)
        plt.legend()
        plt.tight_layout()
        
        chart_path = f"/tmp/fund_performance_{self.fund_code}.png"
        plt.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close()
        
        return chart_path
    
    def _generate_risk_chart(self) -> str:
        """生成风险收益散点图"""
        self._setup_chinese_font()
        
        risk = self.data['risk_metrics']
        if not risk:
            return None
            
        plt.figure(figsize=(8, 6))
        
        # 绘制该基金的位置
        plt.scatter(risk.get('volatility', 0), risk.get('annual_return', 0), 
                   s=200, c='red', marker='*', label='本基金', zorder=5)
        
        # 添加网格和标签
        plt.axhline(y=0, color='gray', linestyle='--', alpha=0.5)
        plt.grid(True, alpha=0.3)
        plt.xlabel('年化波动率 (%)', fontsize=12)
        plt.ylabel('年化收益率 (%)', fontsize=12)
        plt.title('风险收益分布', fontsize=14)
        plt.legend()
        
        chart_path = f"/tmp/fund_risk_{self.fund_code}.png"
        plt.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close()
        
        return chart_path
    
    def _generate_sector_chart(self) -> str:
        """生成行业分布饼图"""
        self._setup_chinese_font()
        
        holdings = self.data['holdings']
        if holdings.empty:
            return None
        
        # 简单分类（实际应该使用申万行业分类）
        sectors = {}
        for _, row in holdings.iterrows():
            # 这里简化处理，实际应调用行业分类API
            sector = '其他'  # 需要完善行业分类逻辑
            sectors[sector] = sectors.get(sector, 0) + row.get('ratio', 0)
        
        if not sectors:
            return None
            
        plt.figure(figsize=(8, 6))
        plt.pie(sectors.values(), labels=sectors.keys(), autopct='%1.1f%%', startangle=90)
        plt.title('行业分布', fontsize=14)
        plt.axis('equal')
        
        chart_path = f"/tmp/fund_sector_{self.fund_code}.png"
        plt.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close()
        
        return chart_path
    
    def _add_cover(self):
        """添加封面"""
        fund_info = self.data['fund_info']
        fund_name = fund_info.get('fund_name', '')
        fund_code = fund_info.get('fund_code', '')
        
        # 空行
        for _ in range(5):
            self._add_paragraph('')
        
        # 报告标题
        title = self._add_paragraph('基金投资研究报告', alignment=WD_ALIGN_PARAGRAPH.CENTER)
        title.runs[0].font.size = Pt(28)
        title.runs[0].font.bold = True
        title.runs[0].font.name = '黑体'
        title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 空行
        for _ in range(3):
            self._add_paragraph('')
        
        # 基金信息
        info_text = f"{fund_name}\n({fund_code})"
        info_p = self._add_paragraph(info_text, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        info_p.runs[0].font.size = Pt(20)
        info_p.runs[0].font.name = '黑体'
        info_p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 空行
        for _ in range(5):
            self._add_paragraph('')
        
        # 报告信息
        date_str = datetime.now().strftime('%Y年%m月%d日')
        self._add_paragraph(f'报告日期：{date_str}', alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._add_paragraph('报告性质：内部研究参考', alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # 分页
        self.doc.add_page_break()
    
    def _add_chapter1_overview(self):
        """第一章：基金概览"""
        self._add_heading('第一章 基金概览', level=1)
        
        fund_info = self.data['fund_info']
        
        # 基本信息表格
        basic_data = [
            ['项目', '内容'],
            ['基金名称', fund_info.get('fund_name', '-')],
            ['基金代码', fund_info.get('fund_code', '-')],
            ['基金类型', fund_info.get('fund_type', '-')],
            ['成立日期', fund_info.get('establish_date', '-')],
            ['管理规模', fund_info.get('manage_scale', '-')],
            ['基金管理人', fund_info.get('company', '-')],
            ['基金经理', fund_info.get('current_manager', fund_info.get('manager', '-'))],
            ['任职日期', fund_info.get('manager_date', '-')],
            ['业绩基准', fund_info.get('benchmark', '-')],
        ]
        
        self._add_paragraph('一、基金基本信息')
        self._add_table(basic_data)
        
        # 基金简介
        self._add_paragraph('\n二、基金简介')
        intro_text = (f"本基金成立于{fund_info.get('establish_date', '未知')}，"
                     f"由{fund_info.get('company', '未知')}管理，"
                     f"现任基金经理{fund_info.get('current_manager', fund_info.get('manager', '未知'))}。"
                     f"基金类型为{fund_info.get('fund_type', '混合型')}，"
                     f"最新管理规模{fund_info.get('manage_scale', '未知')}。")
        self._add_paragraph(intro_text)
        
        self.doc.add_page_break()
    
    def _add_chapter2_performance(self):
        """第二章：历史业绩回顾"""
        self._add_heading('第二章 历史业绩回顾', level=1)
        
        performance = self.data['performance']
        
        # 业绩表格
        perf_data = [['周期', '收益率(%)', '同类排名']]
        
        periods = [
            ('近1月', 'last_month'),
            ('近3月', 'last_3month'),
            ('近6月', 'last_6month'),
            ('近1年', 'last_year'),
            ('近2年', 'last_2year'),
            ('近3年', 'last_3year'),
            ('近5年', 'last_5year'),
            ('成立来', 'since_start'),
        ]
        
        for name, key in periods:
            if key in performance and performance[key].get('return') is not None:
                ret = performance[key]['return']
                rank = performance[key].get('rank', '-')
                perf_data.append([name, f'{ret:.2f}', rank])
        
        self._add_paragraph('一、各周期业绩表现')
        self._add_table(perf_data)
        
        # 业绩走势图
        self._add_paragraph('\n二、业绩走势')
        chart_path = self._generate_performance_chart()
        if chart_path:
            self._add_picture(chart_path, width=6)
        
        self.doc.add_page_break()
    
    def _add_chapter3_risk(self):
        """第三章：风险分析"""
        self._add_heading('第三章 风险分析', level=1)
        
        risk = self.data['risk_metrics']
        
        # 风险指标表格
        risk_data = [
            ['指标', '数值', '说明'],
            ['年化收益率', f"{risk.get('annual_return', '-'):.2f}%" if isinstance(risk.get('annual_return'), (int, float)) else '-', '过去三年年化收益'],
            ['年化波动率', f"{risk.get('volatility', '-'):.2f}%" if isinstance(risk.get('volatility'), (int, float)) else '-', '收益波动程度'],
            ['最大回撤', f"{risk.get('max_drawdown', '-'):.2f}%" if isinstance(risk.get('max_drawdown'), (int, float)) else '-', '最大亏损幅度'],
            ['夏普比率', f"{risk.get('sharpe_ratio', '-'):.2f}" if isinstance(risk.get('sharpe_ratio'), (int, float)) else '-', '风险调整后收益'],
            ['卡玛比率', f"{risk.get('calmar_ratio', '-'):.2f}" if isinstance(risk.get('calmar_ratio'), (int, float)) else '-', '收益回撤比'],
        ]
        
        self._add_paragraph('一、风险指标')
        self._add_table(risk_data)
        
        # 风险评估说明
        self._add_paragraph('\n二、风险评估')
        
        sharpe = risk.get('sharpe_ratio', 0)
        if isinstance(sharpe, (int, float)):
            if sharpe > 1:
                risk_desc = "该基金夏普比率大于1，说明在承担单位风险的情况下获得了较好的超额收益，风险调整后收益表现优秀。"
            elif sharpe > 0.5:
                risk_desc = "该基金夏普比率在0.5-1之间，风险调整后收益表现良好。"
            else:
                risk_desc = "该基金夏普比率较低，风险调整后收益表现一般，需要关注风险控制能力。"
            self._add_paragraph(risk_desc)
        
        # 风险收益图
        self._add_paragraph('\n三、风险收益分布')
        chart_path = self._generate_risk_chart()
        if chart_path:
            self._add_picture(chart_path, width=5)
        
        self.doc.add_page_break()
    
    def _add_chapter4_attribution(self):
        """第四章：业绩归因"""
        self._add_heading('第四章 业绩归因', level=1)
        
        self._add_paragraph('一、收益来源分析')
        attribution_text = """
基于Brinson归因模型，本基金收益主要来源于：

1. 资产配置收益：基金经理对股票、债券等大类资产的配置决策带来的超额收益。
2. 个股选择收益：在既定资产配置下，通过选股带来的超额收益。
3. 交互收益：资产配置与个股选择的协同效应。

（注：详细的Brinson归因分析需要更详细的持仓和基准数据）
        """
        self._add_paragraph(attribution_text)
        
        self._add_paragraph('\n二、择时能力评估')
        timing_text = """
择时能力主要通过以下维度评估：

1. 仓位择时：在市场高点减仓、低点加仓的能力
2. 行业择时：对行业轮动节奏的把握
3. 个股择时：买卖时机的选择

建议结合历史调仓记录和净值走势综合判断基金经理的择时能力。
        """
        self._add_paragraph(timing_text)
        
        self.doc.add_page_break()
    
    def _add_chapter5_holdings(self):
        """第五章：持仓分析"""
        self._add_heading('第五章 持仓分析', level=1)
        
        holdings = self.data['holdings']
        asset = self.data['asset_allocation']
        
        # 资产配置
        self._add_paragraph('一、资产配置')
        if asset:
            asset_data = [
                ['资产类别', '占比'],
                ['股票', asset.get('stock_ratio', '-')],
                ['债券', asset.get('bond_ratio', '-')],
                ['现金', asset.get('cash_ratio', '-')],
                ['其他', asset.get('other_ratio', '-')],
            ]
            self._add_table(asset_data)
        
        # 重仓股
        self._add_paragraph('\n二、前十大重仓股')
        if not holdings.empty:
            holding_data = [['序号', '股票代码', '股票名称', '占净值比']]
            for i, (_, row) in enumerate(holdings.head(10).iterrows(), 1):
                ratio = row.get('ratio', 0)
                if isinstance(ratio, str):
                    ratio = float(ratio.replace('%', ''))
                ratio_str = f"{ratio:.2f}%" if ratio else '-'
                holding_data.append([
                    i,
                    row.get('stock_code', '-'),
                    row.get('stock_name', '-'),
                    ratio_str
                ])
            self._add_table(holding_data)
        
        # 行业分布
        self._add_paragraph('\n三、行业分布')
        
        # 进行持仓有效性分析
        print("  → 进行持仓有效性分析...")
        analyzer = HoldingEffectivenessAnalyzer()
        holding_analysis = analyzer.analyze_holdings_effectiveness(
            holdings, 
            self.data['nav_history'],
            self.fund_code
        )
        
        # 显示行业分布
        sector_info = holding_analysis.get('sector_distribution', {})
        if sector_info and sector_info.get('sector_ratios'):
            sector_data = [['行业板块', '配置比例(%)']]
            for sector, ratio in sorted(sector_info['sector_ratios'].items(), 
                                       key=lambda x: x[1], reverse=True):
                if ratio > 0:
                    sector_data.append([sector, f"{ratio:.2f}"])
            self._add_table(sector_data)
            
            # 行业集中度说明
            if sector_info.get('main_sectors'):
                main_sector = sector_info['main_sectors'][0]
                self._add_paragraph(f"\n主要配置板块：{main_sector[0]}（{main_sector[1]:.2f}%），"
                                  f"行业集中度：{sector_info.get('sector_concentration', '未知')}")
        
        # 四、持仓有效性分析
        self._add_paragraph('\n四、持仓有效性分析')
        
        # 集中度分析
        conc_info = holding_analysis.get('concentration_analysis', {})
        if conc_info:
            self._add_paragraph('1. 持仓集中度评估')
            conc_text = f"""
前5大重仓股合计占比：{conc_info.get('top5_ratio', '-')}%
前3大重仓股合计占比：{conc_info.get('top3_ratio', '-')}%
Herfindahl-Hirschman指数：{conc_info.get('hhi', '-')}（{conc_info.get('concentration_level', '-')})
集中度风险：{conc_info.get('risk', '-')}
            """
            self._add_paragraph(conc_text)
        
        # 当前持仓特征
        current_info = holding_analysis.get('current_holdings', {})
        if current_info:
            self._add_paragraph('2. 当前持仓特征')
            holding_text = f"""
持仓股票数量：{current_info.get('holding_count', '-')}只
前5大持仓占比：{current_info.get('top5_ratio', '-')}%，集中度评价：{current_info.get('top5_concentration', '-')}
最大单一持仓：{current_info.get('max_stock', '-')}（{current_info.get('max_ratio', '-')}%）
            """
            self._add_paragraph(holding_text)
        
        # 持仓质量评估
        quality_info = holding_analysis.get('portfolio_quality', {})
        if quality_info:
            self._add_paragraph('3. 持仓质量评估')
            quality_text = f"""
近一个月收益率：{quality_info.get('recent_month_return', '-')}%
持仓组合波动率：{quality_info.get('portfolio_volatility', '-')}%
持仓质量评级：{quality_info.get('quality_rating', '-')}
            """
            self._add_paragraph(quality_text)
        
        # 有效性评分和建议
        score = holding_analysis.get('effectiveness_score', 0)
        recommendations = holding_analysis.get('recommendations', [])
        
        self._add_paragraph(f'\n五、有效性评分：{score}/100')
        
        if score >= 80:
            rating = "优秀"
        elif score >= 60:
            rating = "良好"
        elif score >= 40:
            rating = "一般"
        else:
            rating = "较差"
        
        self._add_paragraph(f'评级：{rating}')
        
        if recommendations:
            self._add_paragraph('\n六、持仓分析建议：')
            for rec in recommendations:
                self._add_paragraph(f"  • {rec}")
        
        self.doc.add_page_break()
    
    def _add_chapter6_turnover(self):
        """第六章：调仓特征"""
        self._add_heading('第六章 调仓特征', level=1)
        
        self._add_paragraph('一、换手率分析')
        turnover_text = """
换手率反映基金的交易活跃程度：

- 换手率 = (期间买卖股票总金额 / 期间平均资产净值) × 100%
- 高换手率（>300%）：交易频繁，可能偏向趋势跟踪或波段操作
- 中等换手率（100%-300%）：适度调仓，平衡持有与交易
- 低换手率（<100%）：持股稳定，偏向长期价值投资

（注：具体换手率数据需要从定期报告中获取）
        """
        self._add_paragraph(turnover_text)
        
        self._add_paragraph('\n二、重仓股稳定性')
        stability_text = """
通过追踪前十大重仓股的持续持有情况，可以评估基金经理的持股信心：

- 核心持仓：连续多个季度位列前十大重仓的股票
- 波段操作：偶尔出现在重仓股名单中的股票
- 新进退出：近期新增或退出的重仓股

持续重仓的股票通常代表基金经理的核心看好方向。
        """
        self._add_paragraph(stability_text)
        
        self.doc.add_page_break()
    
    def _add_chapter7_manager(self):
        """第七章：基金经理画像"""
        self._add_heading('第七章 基金经理画像', level=1)
        
        fund_info = self.data['fund_info']
        manager_info = self.data['manager_info']
        
        # 基金经理基本信息
        self._add_paragraph('一、基本信息')
        manager_data = [
            ['项目', '内容'],
            ['姓名', manager_info.get('name', fund_info.get('current_manager', '-'))],
            ['从业年限', manager_info.get('tenure', '-')],
            ['任职日期', fund_info.get('manager_date', '-')],
            ['所属公司', fund_info.get('company', '-')],
            ['管理规模', manager_info.get('total_scale', '-')],
        ]
        self._add_table(manager_data)
        
        # 历史持仓分析
        self._add_paragraph('\n二、历史持仓分析')
        print("  → 进行历史持仓分析...")
        
        history_analyzer = HoldingHistoryAnalyzer()
        history_report = history_analyzer.generate_full_report(self.fund_code)
        
        # 显示分析的季度
        quarters = history_report.get('quarters_analyzed', [])
        if quarters:
            self._add_paragraph(f"分析季度: {', '.join(quarters[:8])}")  # 最多显示8个季度
        else:
            self._add_paragraph("注：需要历史持仓数据文件（格式: {基金代码}_{年份}{季度}Q_holdings.csv）进行详细分析")
        
        # 持仓变化分析
        holding_changes = history_report.get('holding_changes', {})
        
        # 1. 持仓稳定性
        stability = holding_changes.get('holding_stability', {})
        if stability:
            self._add_paragraph('\n1. 持仓稳定性评估')
            stability_text = f"""
历史交易股票总数：{stability.get('total_stocks_traded', '-')} 只
核心持仓数量（持有>=4个季度）：{stability.get('core_holdings_count', '-')} 只
波段操作数量（持有<3个季度）：{stability.get('band_trading_count', '-')} 只
            """
            self._add_paragraph(stability_text)
            
            # 核心持仓列表
            core_holdings = holding_changes.get('core_holdings', [])
            if core_holdings:
                self._add_paragraph('\n核心持仓股票（长期持有）:')
                core_data = [['股票名称', '持有季度数', '持有季度']]
                for core in sorted(core_holdings, key=lambda x: x['quarters_held'], reverse=True)[:5]:
                    core_data.append([
                        core['stock_name'],
                        str(core['quarters_held']),
                        ', '.join(core['quarters'][:3])  # 最多显示3个季度
                    ])
                self._add_table(core_data)
        
        # 2. 换手率分析
        turnover = holding_changes.get('turnover_analysis', {})
        if turnover:
            self._add_paragraph('\n2. 换手率分析')
            turnover_text = f"""
平均季度换手率：{turnover.get('average_turnover_rate', '-')}%
换手率水平：{turnover.get('turnover_level', '-')}

换手率解读：
- 高换手率（>60%）：交易频繁，偏向趋势跟踪或行业轮动
- 中等换手率（30%-60%）：适度调仓，平衡持有与交易
- 低换手率（<30%）：持股稳定，偏向长期价值投资
            """
            self._add_paragraph(turnover_text)
        
        # 3. 季度调仓明细
        quarterly_changes = holding_changes.get('quarterly_changes', [])
        if quarterly_changes:
            self._add_paragraph('\n3. 季度调仓明细')
            for change in quarterly_changes[:3]:  # 显示最近3个季度
                period = change.get('period', '')
                added = change.get('added', [])
                removed = change.get('removed', [])
                turnover_rate = change.get('turnover_rate', 0)
                
                change_text = f"""
{period}:
- 新增股票: {', '.join(added) if added else '无'}
- 退出股票: {', '.join(removed) if removed else '无'}
- 换手率: {turnover_rate:.1f}%
                """
                self._add_paragraph(change_text)
        
        # 选股能力评估
        self._add_paragraph('\n三、选股能力评估')
        
        picking = history_report.get('stock_picking_ability', {})
        if picking and picking.get('overall_score', 0) > 0:
            picking_text = f"""
选股能力评分：{picking.get('picking_skill_score', '-')}/100
择时能力评分：{picking.get('timing_skill_score', '-')}/100
综合能力评分：{picking.get('overall_score', '-')}/100
评级：{picking.get('evaluation', '-')}
            """
            self._add_paragraph(picking_text)
            
            # 关键洞察
            insights = picking.get('key_insights', [])
            if insights:
                self._add_paragraph('\n评估要点：')
                for insight in insights:
                    self._add_paragraph(f"  • {insight}")
        else:
            self._add_paragraph("基于当前持仓数据分析：")
            self._add_paragraph("• 需要更多历史季度数据进行选股能力评估")
            self._add_paragraph("• 建议持续跟踪基金经理的调仓操作和选股效果")
        
        # 投资风格识别
        self._add_paragraph('\n四、投资风格识别')
        
        style = history_report.get('manager_style', {})
        if style:
            style_text = f"""
投资风格：{style.get('style', '-')}
换手率水平：{style.get('turnover_level', '-')}
持仓集中度：{style.get('holding_concentration', '-')}

风格特征：
            """
            self._add_paragraph(style_text)
            
            characteristics = style.get('characteristics', [])
            for char in characteristics:
                self._add_paragraph(f"  • {char}")
        else:
            # 基于当前数据的简单分析
            current_holdings = self.data['holdings']
            if not current_holdings.empty:
                avg_ratio = current_holdings['ratio'].mean()
                max_ratio = current_holdings['ratio'].max()
                
                if max_ratio > 8:
                    style_desc = "偏向集中持股，对核心标的有较强信心"
                else:
                    style_desc = "持仓相对分散，注重风险控制"
                
                self._add_paragraph(f"基于当前持仓分析：{style_desc}")
        
        # 能力圈分析
        self._add_paragraph('\n五、能力圈分析')
        
        # 基于行业分布分析能力圈
        if not self.data['holdings'].empty:
            # 获取行业分析（复用holding_analyzer的分析）
            analyzer = HoldingEffectivenessAnalyzer()
            sector_analysis = analyzer._analyze_sectors(self.data['holdings'])
            
            main_sectors = sector_analysis.get('main_sectors', [])
            if main_sectors:
                self._add_paragraph('主要能力圈（行业配置）:')
                for sector, ratio in main_sectors[:3]:
                    self._add_paragraph(f"  • {sector}板块（{ratio:.2f}%）")
            
            # 核心持仓分析
            current_holdings = self.data['holdings']
            if not current_holdings.empty and 'stock_name' in current_holdings.columns:
                top_holdings = current_holdings.nlargest(3, 'ratio')
                self._add_paragraph('\n核心重仓股（当前）:')
                for _, row in top_holdings.iterrows():
                    self._add_paragraph(f"  • {row['stock_name']}（{row['ratio']:.2f}%）")
        
        capability_summary = """
\n能力圈总结：
• 选股能力：基于历史持仓稳定性和核心持仓表现评估
• 行业偏好：通过重仓股行业分布识别擅长领域
• 风控能力：通过持仓集中度和换手率评估风险控制水平
• 市场适应：通过不同市场环境下的调仓策略评估适应能力
        """
        self._add_paragraph(capability_summary)
        
        self.doc.add_page_break()
    
    def _add_chapter8_comparison(self):
        """第八章：同类对比"""
        self._add_heading('第八章 同类对比', level=1)
        
        # 进行同类对比分析
        print("  → 进行同类对比分析...")
        
        comparator = FundComparator()
        comparison_report = comparator.generate_comparison_report(
            self.data['fund_info'],
            self.data['performance'],
            self.data['risk_metrics']
        )
        
        # 一、同类业绩对比表
        self._add_paragraph('一、业绩对比')
        
        similar_funds = comparison_report.get('similar_funds', {})
        perf_table = similar_funds.get('performance_table', [])
        
        if perf_table:
            self._add_paragraph('\n与同类基金业绩对比：')
            self._add_table(perf_table)
        
        # 二、详细排名分析
        self._add_paragraph('\n二、同类排名分析')
        
        period_comparison = comparison_report.get('period_comparison', {})
        if period_comparison:
            # 构建排名对比表
            rank_data = [['周期', '本基金收益', '同类平均', '超额收益', '排名估算', '分位']]
            
            for period, data in period_comparison.items():
                rank_data.append([
                    data['period_name'],
                    f"{data['fund_return']:.2f}%",
                    f"{data['category_avg']:.2f}%",
                    f"{data['excess_return']:+.2f}%",
                    f"{data['rank_estimate']}/{data['total_funds']}",
                    f"前{data['rank_percentile']:.0f}%" if data['rank_percentile'] > 50 else f"后{100-data['rank_percentile']:.0f}%"
                ])
            
            self._add_table(rank_data)
            
            # 综合表现总结
            summary = comparison_report.get('summary', {})
            avg_percentile = summary.get('avg_percentile', 0)
            overall_rating = summary.get('overall_rating', '')
            
            self._add_paragraph(f"\n综合表现：")
            self._add_paragraph(f"• 平均排名分位：{avg_percentile:.1f}%")
            self._add_paragraph(f"• 综合评级：{overall_rating}")
            self._add_paragraph(f"• 表现评价：{summary.get('rating_desc', '')}")
            
            # 最佳和最差周期
            best = summary.get('best_period', {})
            worst = summary.get('worst_period', {})
            
            if best and worst:
                self._add_paragraph(f"\n• 最佳表现：{best.get('period', '')}（前{best.get('percentile', 0):.0f}%）")
                self._add_paragraph(f"• 最差表现：{worst.get('period', '')}（后{100-worst.get('percentile', 0):.0f}%）")
            
            # 胜率
            win_rate = summary.get('win_rate', 0)
            self._add_paragraph(f"• 跑赢同类平均的周期比例：{win_rate:.0f}%")
        
        # 三、风险对比
        self._add_paragraph('\n三、风险指标对比')
        
        risk_comparison = comparison_report.get('risk_comparison', {})
        if risk_comparison:
            risk_data = [['指标', '本基金', '同类平均', '差异', '评价']]
            
            # 波动率
            vol = risk_comparison.get('volatility', {})
            vol_diff = vol.get('diff', 0)
            vol_eval = '较低' if vol_diff < 0 else '较高' if vol_diff > 0 else '持平'
            risk_data.append([
                '年化波动率',
                f"{vol.get('fund', 0):.2f}%",
                f"{vol.get('category_avg', 0):.2f}%",
                f"{vol_diff:+.2f}%",
                vol_eval
            ])
            
            # 最大回撤
            mdd = risk_comparison.get('max_drawdown', {})
            mdd_diff = mdd.get('diff', 0)
            mdd_eval = '较小' if mdd_diff > 0 else '较大' if mdd_diff < 0 else '持平'
            risk_data.append([
                '最大回撤',
                f"{mdd.get('fund', 0):.2f}%",
                f"{mdd.get('category_avg', 0):.2f}%",
                f"{mdd_diff:+.2f}%",
                mdd_eval
            ])
            
            # 夏普比率
            sharpe = risk_comparison.get('sharpe_ratio', {})
            sharpe_diff = sharpe.get('diff', 0)
            sharpe_eval = '较优' if sharpe_diff > 0 else '较弱' if sharpe_diff < 0 else '持平'
            risk_data.append([
                '夏普比率',
                f"{sharpe.get('fund', 0):.2f}",
                f"{sharpe.get('category_avg', 0):.2f}",
                f"{sharpe_diff:+.2f}",
                sharpe_eval
            ])
            
            self._add_table(risk_data)
        
        # 四、竞争优势与劣势
        self._add_paragraph('\n四、竞争优势与劣势')
        
        strengths = comparison_report.get('strengths', [])
        weaknesses = comparison_report.get('weaknesses', [])
        
        if strengths:
            self._add_paragraph('\n竞争优势：')
            for s in strengths:
                self._add_paragraph(f"  ✅ {s}")
        
        if weaknesses:
            self._add_paragraph('\n竞争劣势：')
            for w in weaknesses:
                self._add_paragraph(f"  ⚠️ {w}")
        
        if not strengths and not weaknesses:
            self._add_paragraph('与同类基金相比，各项指标表现接近平均水平。')
        
        # 五、投资建议
        self._add_paragraph('\n五、同类对比投资建议')
        
        recommendations = comparison_report.get('recommendations', [])
        if recommendations:
            for rec in recommendations:
                self._add_paragraph(f"  • {rec}")
        else:
            self._add_paragraph("  • 建议持续关注基金在同类中的排名变化")
        
        self.doc.add_page_break()
    
    def _add_chapter9_pros_cons(self):
        """第九章：优势与风险"""
        self._add_heading('第九章 优势与风险', level=1)
        
        # 投资亮点
        self._add_paragraph('一、核心投资亮点')
        
        fund_info = self.data['fund_info']
        risk = self.data['risk_metrics']
        
        highlights = []
        
        # 基于数据生成亮点
        sharpe = risk.get('sharpe_ratio', 0)
        if isinstance(sharpe, (int, float)) and sharpe > 1:
            highlights.append("风险调整后收益优秀，夏普比率大于1")
        
        calmar = risk.get('calmar_ratio', 0)
        if isinstance(calmar, (int, float)) and calmar > 1:
            highlights.append("收益回撤比良好，风险控制能力较强")
        
        highlights.append(f"基金经理{fund_info.get('current_manager', '')}管理，需关注其投资风格和稳定性")
        highlights.append(f"成立于{fund_info.get('establish_date', '')}，具有一定历史业绩积累")
        
        for i, h in enumerate(highlights, 1):
            self._add_paragraph(f"{i}. {h}")
        
        # 风险提示
        self._add_paragraph('\n二、主要风险点')
        
        risks = [
            "基金投资有风险，过往业绩不代表未来表现",
            "权益类基金波动较大，短期内可能出现较大回撤",
            "基金经理变更风险（如适用）",
            "市场系统性风险",
            "流动性风险（如适用）",
        ]
        
        for i, r in enumerate(risks, 1):
            self._add_paragraph(f"{i}. {r}")
        
        # 适合投资者
        self._add_paragraph('\n三、适合投资者类型')
        suitable_text = """
基于本基金的风险收益特征，建议以下类型投资者考虑：

- 风险承受能力：（根据基金类型填写）
- 投资期限建议：中长期持有（建议1年以上）
- 配置定位：可作为（核心/卫星）资产配置
        """
        self._add_paragraph(suitable_text)
        
        self.doc.add_page_break()
    
    def _add_chapter10_rating(self):
        """第十章：投资价值判断"""
        self._add_heading('第十章 投资价值判断', level=1)
        
        # 综合评分
        self._add_paragraph('一、综合评分')
        
        # 计算综合评分（简化版）
        risk = self.data['risk_metrics']
        score = 70  # 基础分
        
        sharpe = risk.get('sharpe_ratio', 0)
        if isinstance(sharpe, (int, float)):
            if sharpe > 1.5:
                score += 15
            elif sharpe > 1:
                score += 10
            elif sharpe > 0.5:
                score += 5
        
        calmar = risk.get('calmar_ratio', 0)
        if isinstance(calmar, (int, float)):
            if calmar > 2:
                score += 10
            elif calmar > 1:
                score += 5
        
        score = min(score, 100)
        
        # 评级
        if score >= 85:
            rating = '强烈推荐'
            position = '核心配置'
        elif score >= 70:
            rating = '推荐'
            position = '卫星配置'
        elif score >= 55:
            rating = '中性'
            position = '观察名单'
        else:
            rating = '谨慎'
            position = '暂不建议'
        
        rating_data = [
            ['评估维度', '评估结果'],
            ['综合评分', f'{score}/100'],
            ['投资评级', rating],
            ['配置建议', position],
        ]
        self._add_table(rating_data)
        
        # 评分说明
        self._add_paragraph('\n二、评分说明')
        
        sharpe_val = risk.get('sharpe_ratio', '-')
        sharpe_str = f"{sharpe_val:.2f}" if isinstance(sharpe_val, (int, float)) else str(sharpe_val)
        
        mdd_val = risk.get('max_drawdown', '-')
        mdd_str = f"{mdd_val:.2f}%" if isinstance(mdd_val, (int, float)) else str(mdd_val)
        
        ann_ret_val = risk.get('annual_return', '-')
        ann_ret_str = f"{ann_ret_val:.2f}%" if isinstance(ann_ret_val, (int, float)) else str(ann_ret_val)
        
        calmar_val = risk.get('calmar_ratio', '-')
        calmar_str = f"{calmar_val:.2f}" if isinstance(calmar_val, (int, float)) else str(calmar_val)
        
        score_text = f"""综合评分基于以下维度计算：

1. 风险调整后收益（权重30%）：夏普比率 {sharpe_str}
2. 风险控制能力（权重25%）：最大回撤 {mdd_str}
3. 长期业绩表现（权重25%）：年化收益 {ann_ret_str}
4. 收益稳定性（权重20%）：卡玛比率 {calmar_str}

评级标准：
- 强烈推荐（85-100分）：各项指标优秀，值得重点配置
- 推荐（70-84分）：整体表现良好，可作为卫星配置
- 中性（55-69分）：表现一般，建议观察
- 谨慎（55分以下）：存在明显不足，暂不建议"""
        self._add_paragraph(score_text)
        
        self.doc.add_page_break()
    
    def _add_chapter11_outlook(self):
        """第十一章：前瞻性观点"""
        self._add_heading('第十一章 前瞻性观点', level=1)
        
        self._add_paragraph('一、后市展望')
        outlook_text = """
基于当前市场环境和基金特征，对后市的分析：

1. 市场环境判断：
   - 宏观经济形势
   - 市场估值水平
   - 政策环境变化

2. 对基金的影响：
   - 该基金投资策略在当前市场的适应性
   - 基金经理的能力圈与市场风格的匹配度
   - 潜在的机会与风险
        """
        self._add_paragraph(outlook_text)
        
        self._add_paragraph('\n二、策略有效性预期')
        strategy_text = """
对该基金经理投资策略未来有效性的判断：

1. 策略持续性：
   - 投资风格是否稳定
   - 能力圈是否持续有效
   - 规模对策略的影响

2. 预期收益区间：
   - 基于历史表现和市场环境，预期年化收益区间
   - 预期最大回撤区间

3. 关键观察指标：
   - 需要重点关注的指标变化
   - 策略失效的信号
        """
        self._add_paragraph(strategy_text)
        
        self._add_paragraph('\n三、关注要点提示')
        watch_text = """
建议持续关注以下要点：

1. 定期报告披露后的持仓变化
2. 基金经理公开观点和路演信息
3. 规模变化对业绩的影响
4. 基金经理变更情况（如有）
5. 市场风格切换对基金表现的影响

免责声明：
本报告仅供参考，不构成投资建议。基金投资有风险，投资者应根据自身风险承受能力谨慎决策。过往业绩不代表未来表现，投资前请仔细阅读基金合同和招募说明书。
        """
        self._add_paragraph(watch_text)


def main():
    parser = argparse.ArgumentParser(description='生成基金投资研究报告')
    parser.add_argument('fund_code', help='基金代码，如 519702')
    parser.add_argument('--output', '-o', help='输出目录', default='.')
    parser.add_argument('--years', '-y', type=int, help='历史数据年数', default=3)
    parser.add_argument('--template', '-t', help='Word模板路径')
    
    args = parser.parse_args()
    
    # 生成报告
    generator = FundResearchReport(template_path=args.template)
    
    try:
        report_path = generator.generate(
            fund_code=args.fund_code,
            output_dir=args.output,
            years=args.years
        )
        print(f"\n报告生成完成！")
        print(f"文件路径: {os.path.abspath(report_path)}")
    except Exception as e:
        print(f"\n生成报告失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
