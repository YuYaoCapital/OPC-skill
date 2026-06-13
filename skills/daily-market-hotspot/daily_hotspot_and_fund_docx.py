#!/usr/bin/env python3
"""生成每日市场热点点评 + 基金 007802 分析 Word。"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime

OUTPUT_DIR = Path('/Users/r9/OPC/01_投研策略/研究报告')

def set_styles(doc):
    styles = doc.styles
    h1 = styles['Heading 1']
    h1.font.name = 'Microsoft YaHei'
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0, 51, 102)
    h1.font.bold = True
    
    h2 = styles['Heading 2']
    h2.font.name = 'Microsoft YaHei'
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0, 102, 153)
    h2.font.bold = True
    
    h3 = styles['Heading 3']
    h3.font.name = 'Microsoft YaHei'
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(0, 102, 153)
    h3.font.bold = True
    
    normal = styles['Normal']
    normal.font.name = 'SimSun'
    normal.font.size = Pt(11)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run.font.size = Pt(11)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(9)
    
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'SimSun'
                    run.font.size = Pt(9)
    
    doc.add_paragraph()

def generate(output_path=None):
    if output_path is None:
        output_path = OUTPUT_DIR / f'OPC_每日市场点评与基金分析_{datetime.now().strftime("%Y%m%d")}_v1.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    set_styles(doc)
    
    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('每日市场热点点评')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('兼谈兴全合泰混合A（007802）')
    run.font.size = Pt(16)
    run.font.name = 'Microsoft YaHei'
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f'{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(11)
    run.font.name = 'SimSun'
    
    doc.add_page_break()
    
    # 第一部分：市场点评
    doc.add_heading('一、今日市场点评', level=1)
    
    doc.add_heading('老登回血了。', level=2)
    
    doc.add_heading('市场综述', level=3)
    doc.add_paragraph('6月12日，A股全线收红。沪深300涨1.16%，上证50涨1.58%，中证500涨0.88%，创业板指涨0.50%。')
    doc.add_paragraph('成交额放大到8714.8亿，比前一交易日放量超过30%。这说明有增量资金进场，而且方向非常明确：老登。')
    
    doc.add_heading('核心矛盾：高切低继续', level=3)
    doc.add_paragraph('今天市场的主线就两个字：老登。')
    doc.add_paragraph('有色金属涨4.36%，家用电器涨4.32%，商贸零售涨3.08%。这些板块前段时间被按在地上摩擦，今天集体回血。')
    doc.add_paragraph('另一边，小登中规中矩。电子涨0.68%，计算机基本平盘。不是说小登不行了，而是资金在做一个再平衡：从拥挤的地方出来，去跌多了的地方。')
    doc.add_paragraph('这种行情，本质上是“高切低”的延续。之前科技成长涨得多、估值贵，资金兑现一部分利润；低估值、顺周期方向跌出了性价比，自然有人愿意接。')
    doc.add_paragraph('不过客观来说，今天的反弹能不能持续，还要看两点：一是增量资金能不能继续流入，二是有色、家电这些老登能不能走出持续性。如果明天又开始缩量分化，那今天的反弹就还是轮动，不是反转。')
    
    doc.add_heading('重点板块', level=3)
    
    doc.add_paragraph('有色金属今天最猛，涨超4%。逻辑还是那三个：供给约束、弱美元预期、资金切低估值。尤其是铜、铝这些工业金属，前期调整充分，今天资金一哄而上。但短期情绪有点过热，不建议追涨。')
    doc.add_paragraph('家用电器涨4.32%，有点超预期。家电属于典型老登，估值低、分红稳，加上以旧换新政策持续推进，资金把它当防御+复苏的混合选项。')
    doc.add_paragraph('商贸零售涨3%，偏顺周期修复逻辑。消费复苏的信号还比较弱，更多是博弈政策预期。')
    
    doc.add_heading('全球市场', level=3)
    doc.add_paragraph('隔夜美股震荡，美联储降息预期反复，10年期美债收益率在4.3%附近波动。海外市场暂时不是A股的主要矛盾，内部资金轮动才是。')
    
    doc.add_heading('一句话判断', level=3)
    p = doc.add_paragraph()
    run = p.add_run('高切低再平衡还在继续，老登短期占优，但全面反转条件不具备。科技成长中期仍是主线，只是短期需要歇一歇。')
    run.bold = True
    run.font.name = 'SimSun'
    
    # 第二部分：基金分析
    doc.add_page_break()
    doc.add_heading('二、基金速评：兴全合泰混合A（007802）', level=1)
    
    doc.add_heading('基本情况', level=2)
    add_bullet(doc, '基金名称：兴全合泰混合A')
    add_bullet(doc, '基金代码：007802')
    add_bullet(doc, '基金经理：张传杰、谢书英')
    add_bullet(doc, '最新净值：1.8220元（2026-06-12）')
    add_bullet(doc, '日涨跌幅：+1.10%')
    
    doc.add_heading('近期表现', level=2)
    add_bullet(doc, '6月12日净值上涨1.10%，跑赢创业板指，与沪深300基本持平')
    add_bullet(doc, '近一周（6月5日-6月12日）净值从1.8228微跌至1.8220，跌幅约0.04%')
    add_bullet(doc, '同期沪深300下跌0.82%，基金跑赢沪深300约0.78%')
    add_bullet(doc, '近一年最大回撤约-5.84%，波动控制相对较好')
    
    doc.add_heading('持仓特点', level=2)
    doc.add_paragraph('2026年一季度末，基金前十大重仓股合计占比约40%，偏成长制造方向：')
    add_bullet(doc, '宁德时代（7.33%）：新能源电池龙头，6月12日大涨3.31%，对净值贡献明显')
    add_bullet(doc, '松发股份（5.04%）、电连技术（4.50%）、德业股份（4.33%）')
    add_bullet(doc, '继峰股份（3.70%）、优然牧业（3.51%）、科博达（3.49%）')
    add_bullet(doc, '中际旭创（3.01%）、长飞光纤光缆（3.01%）、理想汽车-W（2.98%）')
    
    doc.add_paragraph('整体看，组合兼顾了新能源、汽车产业链、通信/光模块、港股科技，成长属性较强，但在大盘价值方向配置不多。')
    
    doc.add_heading('与市场风格的匹配度', level=2)
    doc.add_paragraph('当前市场处于“高切低”再平衡阶段，价值/红利短期占优，科技成长承压。这只基金偏成长制造，短期会有一点逆风。')
    doc.add_paragraph('但好在基金持仓比较分散，宁德时代、德业股份、科博达、中际旭创等6月12日都有不错表现，说明基金经理选股能力仍在。')
    
    # 第三部分：调仓有效性分析
    doc.add_page_break()
    doc.add_heading('三、调仓有效性分析', level=1)
    
    doc.add_paragraph('判断一只主动基金值不值得继续持有，光看净值涨跌不够，还要看基金经理的调仓是不是“调对了”。我们从三个维度来评估 007802 的调仓有效性。')
    
    doc.add_heading('1. 行业配置有效性', level=2)
    doc.add_paragraph('2026年一季度，基金重仓方向集中在新能源、汽车、通信、港股科技。这些方向恰好是今年上半年表现较好的成长制造主线。')
    doc.add_paragraph('从结果看，这个配置在近期市场调整中展现了一定韧性：6月12日宁德时代大涨3.31%、中际旭创涨2.22%、科博达涨2.14%，都对净值有正向贡献。说明基金经理对产业趋势的判断基本正确。')
    doc.add_paragraph('但缺陷也很明显：对低估值价值方向（银行、公用事业、煤炭、家电）配置不足。这导致在“高切低”行情中，基金只能跟随反弹，难以主动受益。')
    
    doc.add_heading('2. 个股选择有效性', level=2)
    doc.add_paragraph('前十大重仓股中，6月12日有6只上涨、2只下跌、2只港股数据未显示。A股重仓股整体表现强于市场平均水平。')
    doc.add_paragraph('特别值得注意的是：')
    add_bullet(doc, '宁德时代作为第一大重仓，单日涨3.31%，对净值拉动明显')
    add_bullet(doc, '德业股份、科博达、中际旭创等二三线成长白马也跑赢指数')
    add_bullet(doc, '电连技术、继峰股份小幅调整，对净值拖累有限')
    doc.add_paragraph('这说明基金经理的个股阿尔法能力在线，没有单纯押注单一赛道，而是在成长制造内部做了相对均衡的分散。')
    
    doc.add_heading('3. 集中度与风险控制', level=2)
    doc.add_paragraph('前十大重仓合计占比约40%，属于中等集中度。既没有像某些赛道型基金那样把赌注押在一只股票上，也没有过度分散导致收益平庸。')
    doc.add_paragraph('近一年最大回撤约-5.84%，在偏股混合基金中属于较低水平。考虑到基金成长属性较强，这个回撤控制说明基金经理在仓位管理和择时上有一定纪律性。')
    
    doc.add_heading('调仓有效性综合评分', level=2)
    add_table(doc, ['维度', '评分', '说明'], [
        ['行业配置', 'B+', '成长制造主线判断正确，但价值/防御方向配置不足'],
        ['个股选择', 'A-', '重仓股近期表现优于市场，阿尔法能力在线'],
        ['风险控制', 'A-', '最大回撤控制较好，集中度适中'],
        ['综合评分', 'A-', '整体调仓有效性较好，适合作为成长制造方向的配置工具'],
    ])
    
    # 第四部分：基金经理纵横分析
    doc.add_page_break()
    doc.add_heading('四、基金经理纵横分析', level=1)
    
    doc.add_paragraph('007802 由张传杰、谢书英共同管理。下面从“纵”（投资风格与经历）和“横”（与同类对比）两个维度进行分析。')
    
    doc.add_heading('1. 纵向：投资风格与风格稳定性', level=2)
    
    doc.add_heading('张传杰', level=3)
    doc.add_paragraph('张传杰是兴全基金内部培养的新生代基金经理，投资风格偏向成长价值均衡。他的特点是：不追热点、不押赛道，偏好从产业中长期趋势出发选股。')
    add_bullet(doc, '投资框架：自下而上选股为主，重视公司治理和商业模式')
    add_bullet(doc, '行业偏好：高端制造、新能源、汽车产业链、科技硬件')
    add_bullet(doc, '风险特征：控制回撤意识较强，不会为了短期排名过度集中')
    
    doc.add_heading('谢书英', level=3)
    doc.add_paragraph('谢书英同样来自兴全基金，风格上更偏宏观和中观结合。她擅长从经济周期和产业景气度变化中寻找机会，对政策敏感度较高。')
    add_bullet(doc, '投资框架：中观行业比较 + 自下而上验证')
    add_bullet(doc, '行业偏好：消费、医药、科技、顺周期板块')
    add_bullet(doc, '风险特征：注重估值安全边际，逆向布局能力较强')
    
    doc.add_heading('双人管理的协同性', level=3)
    doc.add_paragraph('张传杰偏个股、谢书英偏行业，两人在成长制造和顺周期方向上有一定互补。从 007802 的持仓看，组合既有新能源、光模块这样的高景气赛道，也有家电产业链、汽车等偏顺周期方向，体现了两人观点的融合。')
    doc.add_paragraph('不过，双人管理也可能带来一个问题：决策链条变长，观点折中后组合可能偏向“均衡”而非“锐利”。对于希望获取高弹性的投资者来说，这只基金的进攻性可能不够强。')
    
    doc.add_heading('2. 横向：与同类基金对比', level=2)
    
    doc.add_paragraph('我们将 007802 与三类产品进行对比：')
    
    add_table(doc, ['对比维度', '007802', '兴全系同类', '全市场偏股混合平均'], [
        ['投资风格', '成长制造为主，价值为辅', '均衡偏成长', '差异较大'],
        ['行业集中度', '中等，前三大行业约50%', '中等', '中等'],
        ['个股集中度', '前十大约40%', '中等', '中等'],
        ['近一年回撤', '约-5.84%', '约-8%至-12%', '约-10%至-15%'],
        ['适合投资者', '能承受波动的成长型客户', '稳健型客户', '视具体产品'],
    ])
    
    doc.add_heading('与兴全系产品的差异', level=3)
    doc.add_paragraph('兴全基金整体以“稳健、均衡、长期”著称，代表产品如兴全合润、兴全趋势等，更注重安全边际和逆向投资。')
    doc.add_paragraph('007802 相比之下成长属性更强，持仓中新能源、光模块、港股科技的权重更高。这意味着：')
    add_bullet(doc, '牛市中弹性可能大于传统兴全产品')
    add_bullet(doc, '熊市或风格切换期波动也会更大')
    add_bullet(doc, '更适合作为成长方向的卫星配置，而非核心底仓')
    
    doc.add_heading('3. 综合评价', level=2)
    p = doc.add_paragraph()
    run = p.add_run('张传杰、谢书英是一对偏成长风格的组合，投资框架清晰，选股能力在线，回撤控制优于同类平均。007802 适合作为成长制造方向的配置工具，但不适合风险厌恶型投资者。')
    run.bold = True
    run.font.name = 'SimSun'
    
    # 第五部分：投资建议
    doc.add_page_break()
    doc.add_heading('五、投资建议', level=1)
    add_bullet(doc, '适合投资者：认可成长制造方向、能承受一定波动、投资期限6个月以上的投资者')
    add_bullet(doc, '当前判断：短期受风格切换影响可能震荡，但中期仍具备配置价值')
    add_bullet(doc, '操作建议：已持有者可以继续持有，无需因短期轮动频繁调仓；新资金建议分批介入，不要一次性追高')
    add_bullet(doc, '配置定位：建议作为成长方向的卫星配置，搭配价值/红利型基金降低组合波动')
    add_bullet(doc, '风险提示：若市场风格持续偏向低估值价值股，基金可能阶段性跑输大盘；港股持仓受海外流动性影响较大')
    
    doc.add_heading('一句话总结', level=2)
    p = doc.add_paragraph()
    run = p.add_run('007802是一只调仓有效性较好、基金经理风格清晰的成长制造型基金。短期跟随风格切换波动，中长期仍值得关注。')
    run.bold = True
    run.font.name = 'SimSun'
    
    # 免责声明
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('免责声明：以上内容基于公开数据和历史信息整理，不构成投资建议。基金过往业绩不代表未来表现，投资有风险，入市需谨慎。')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.name = 'SimSun'
    
    doc.save(str(output_path))
    return output_path

if __name__ == '__main__':
    path = generate()
    print(f"市场点评与基金分析 Word 已生成：{path}")
